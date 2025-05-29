"""
Document Processing Pipeline
Handles uploading and processing various document types for the Knowledge Worker Agent
"""
import os
import hashlib
import logging
from typing import List, Dict, Any, Optional, BinaryIO
from datetime import datetime
import uuid
import mimetypes

# Document processing libraries
import PyPDF2
from docx import Document
from bs4 import BeautifulSoup
import requests

# Azure services
from azure.storage.blob import BlobServiceClient
from azure_search_service import AzureSearchService

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class DocumentProcessor:
    """Handles document processing and indexing for the knowledge worker agent"""
    
    def __init__(self):
        """Initialize the document processor"""
        # Initialize Azure Storage client
        storage_account_name = os.getenv("AZURE_STORAGE_ACCOUNT_NAME")
        storage_account_key = os.getenv("AZURE_STORAGE_ACCOUNT_KEY")
        self.container_name = os.getenv("AZURE_STORAGE_CONTAINER_NAME", "documents")
        
        if storage_account_name and storage_account_key:
            self.blob_service_client = BlobServiceClient(
                account_url=f"https://{storage_account_name}.blob.core.windows.net",
                credential=storage_account_key
            )
        else:
            self.blob_service_client = None
            logger.warning("Azure Storage not configured - file upload disabled")
        
        # Initialize search service
        self.search_service = AzureSearchService()
        
        # Supported file types
        self.supported_types = {
            '.pdf': self._process_pdf,
            '.docx': self._process_docx,
            '.txt': self._process_text,
            '.html': self._process_html,
            '.htm': self._process_html
        }
    
    def _generate_document_id(self, filename: str, content: str) -> str:
        """Generate a unique document ID based on filename and content"""
        content_hash = hashlib.md5(content.encode()).hexdigest()[:8]
        return f"{filename}_{content_hash}"
    
    def _process_pdf(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """Extract text from PDF file"""
        try:
            import io
            pdf_file = io.BytesIO(file_content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            text_content = ""
            for page in pdf_reader.pages:
                text_content += page.extract_text() + "\n"
            
            return {
                "success": True,
                "content": text_content.strip(),
                "page_count": len(pdf_reader.pages),
                "metadata": {
                    "pages": len(pdf_reader.pages),
                    "file_type": "pdf"
                }
            }
        except Exception as e:
            logger.error(f"Error processing PDF {filename}: {str(e)}")
            return {
                "success": False,
                "error": str(e)
            }
    
    def _process_docx(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """Extract text from Word document"""
        try:
            import io
            docx_file = io.BytesIO(file_content)
            doc = Document(docx_file)
            
            text_content = ""
            for paragraph in doc.paragraphs:
                text_content += paragraph.text + "\n"
            
            return {
                "success": True,
                "content": text_content.strip(),
                "paragraph_count": len(doc.paragraphs),
                "metadata": {
                    "paragraphs": len(doc.paragraphs),
                    "file_type": "docx"
                }
            }
        except Exception as e:
            logger.error(f"Error processing DOCX {filename}: {str(e)}")
            return {
                "success": False,
                "error": str(e)
            }
    
    def _process_text(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """Process plain text file"""
        try:
            text_content = file_content.decode('utf-8')
            
            return {
                "success": True,
                "content": text_content.strip(),
                "character_count": len(text_content),
                "metadata": {
                    "characters": len(text_content),
                    "lines": len(text_content.split('\n')),
                    "file_type": "text"
                }
            }
        except Exception as e:
            logger.error(f"Error processing text file {filename}: {str(e)}")
            return {
                "success": False,
                "error": str(e)
            }
    
    def _process_html(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """Extract text from HTML file"""
        try:
            html_content = file_content.decode('utf-8')
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()
            
            # Get text content
            text_content = soup.get_text()
            
            # Clean up whitespace
            lines = (line.strip() for line in text_content.splitlines())
            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
            text_content = ' '.join(chunk for chunk in chunks if chunk)
            
            return {
                "success": True,
                "content": text_content,
                "metadata": {
                    "title": soup.title.string if soup.title else None,
                    "file_type": "html"
                }
            }
        except Exception as e:
            logger.error(f"Error processing HTML {filename}: {str(e)}")
            return {
                "success": False,
                "error": str(e)
            }
    
    def upload_file(self, file_content: bytes, filename: str) -> Optional[str]:
        """Upload file to Azure Blob Storage"""
        try:
            if not self.blob_service_client:
                logger.warning("Azure Storage not configured - skipping file upload")
                return None
            
            blob_name = f"{datetime.now().strftime('%Y/%m/%d')}/{filename}"
            blob_client = self.blob_service_client.get_blob_client(
                container=self.container_name,
                blob=blob_name
            )
            
            blob_client.upload_blob(file_content, overwrite=True)
            logger.info(f"File uploaded to blob storage: {blob_name}")
            
            return f"https://{os.getenv('AZURE_STORAGE_ACCOUNT_NAME')}.blob.core.windows.net/{self.container_name}/{blob_name}"
            
        except Exception as e:
            logger.error(f"Error uploading file to blob storage: {str(e)}")
            return None
    
    def process_file(self, file_content: bytes, filename: str, category: str = "general") -> Dict[str, Any]:
        """Process a file and index it in Azure Search"""
        try:
            # Determine file type
            file_extension = os.path.splitext(filename)[1].lower()
            
            if file_extension not in self.supported_types:
                return {
                    "success": False,
                    "error": f"Unsupported file type: {file_extension}",
                    "supported_types": list(self.supported_types.keys())
                }
            
            # Process the file content
            processor = self.supported_types[file_extension]
            processing_result = processor(file_content, filename)
            
            if not processing_result["success"]:
                return processing_result
            
            content = processing_result["content"]
            
            if not content.strip():
                return {
                    "success": False,
                    "error": "No extractable text content found in file"
                }
            
            # Upload file to blob storage
            blob_url = self.upload_file(file_content, filename)
            
            # Generate document ID
            document_id = self._generate_document_id(filename, content)
            
            # Prepare document for indexing
            document = {
                "id": document_id,
                "title": os.path.splitext(filename)[0],
                "content": content,
                "source": blob_url or filename,
                "category": category,
                "created_date": datetime.now().isoformat(),
                "metadata": {
                    "original_filename": filename,
                    "file_size": len(file_content),
                    "content_length": len(content),
                    **processing_result.get("metadata", {})
                }
            }
            
            # Index the document
            indexing_success = self.search_service.index_document(document)
            
            if indexing_success:
                return {
                    "success": True,
                    "document_id": document_id,
                    "filename": filename,
                    "content_length": len(content),
                    "blob_url": blob_url,
                    "category": category,
                    "processing_details": processing_result.get("metadata", {})
                }
            else:
                return {
                    "success": False,
                    "error": "Failed to index document in search service"
                }
                
        except Exception as e:
            logger.error(f"Error processing file {filename}: {str(e)}")
            return {
                "success": False,
                "error": str(e)
            }
    
    def process_url(self, url: str, category: str = "web") -> Dict[str, Any]:
        """Process content from a URL"""
        try:
            # Fetch the web page
            response = requests.get(url, timeout=30)
            response.raise_for_status()
            
            # Process as HTML
            processing_result = self._process_html(response.content, url)
            
            if not processing_result["success"]:
                return processing_result
            
            content = processing_result["content"]
            
            if not content.strip():
                return {
                    "success": False,
                    "error": "No extractable text content found at URL"
                }
            
            # Generate document ID
            document_id = self._generate_document_id(url, content)
            
            # Prepare document for indexing
            document = {
                "id": document_id,
                "title": processing_result.get("metadata", {}).get("title") or url,
                "content": content,
                "source": url,
                "category": category,
                "created_date": datetime.now().isoformat(),
                "metadata": {
                    "url": url,
                    "content_length": len(content),
                    "fetch_date": datetime.now().isoformat(),
                    **processing_result.get("metadata", {})
                }
            }
            
            # Index the document
            indexing_success = self.search_service.index_document(document)
            
            if indexing_success:
                return {
                    "success": True,
                    "document_id": document_id,
                    "url": url,
                    "title": document["title"],
                    "content_length": len(content),
                    "category": category
                }
            else:
                return {
                    "success": False,
                    "error": "Failed to index document in search service"
                }
                
        except Exception as e:
            logger.error(f"Error processing URL {url}: {str(e)}")
            return {
                "success": False,
                "error": str(e)
            }
    
    def batch_process_files(self, files: List[Dict[str, Any]]) -> Dict[str, Any]:
        """Process multiple files in batch"""
        results = {
            "successful": [],
            "failed": [],
            "total_processed": 0,
            "total_files": len(files)
        }
        
        for file_info in files:
            try:
                result = self.process_file(
                    file_content=file_info["content"],
                    filename=file_info["filename"],
                    category=file_info.get("category", "general")
                )
                
                results["total_processed"] += 1
                
                if result["success"]:
                    results["successful"].append({
                        "filename": file_info["filename"],
                        "document_id": result["document_id"]
                    })
                else:
                    results["failed"].append({
                        "filename": file_info["filename"],
                        "error": result["error"]
                    })
                    
            except Exception as e:
                results["failed"].append({
                    "filename": file_info.get("filename", "unknown"),
                    "error": str(e)
                })
        
        return results
    
    def get_processing_statistics(self) -> Dict[str, Any]:
        """Get statistics about processed documents"""
        try:
            search_stats = self.search_service.get_index_statistics()
            
            return {
                "total_documents": search_stats.get("document_count", 0),
                "index_size": search_stats.get("storage_size", 0),
                "supported_file_types": list(self.supported_types.keys()),
                "storage_configured": self.blob_service_client is not None
            }
        except Exception as e:
            logger.error(f"Error getting processing statistics: {str(e)}")
            return {
                "error": str(e)
            }
